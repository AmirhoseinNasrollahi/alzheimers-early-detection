# paper/build_tables.py
"""
Build paper tables from result CSVs.

Reads:
  - performance_summary.csv
  - confusion_matrices_aggregated.csv

Writes (to --out_dir/tables):
  - Table2_nested_cv_auc.csv
  - Table2_nested_cv_auc.docx
  - Table3_best_config_metrics.csv
  - Table3_best_config_metrics.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path
from typing import Optional

import numpy as np
import pandas as pd


TASK_ORDER = ["CN_vs_AD", "CN_vs_EMCI", "EMCI_vs_AD"]


def _nice_task_label(task: str) -> str:
    return task.replace("_vs_", " vs ")


def _fmt_auc_ci(mean: float, lo: float, hi: float, ndp: int = 3) -> str:
    if np.isnan(mean) or np.isnan(lo) or np.isnan(hi):
        return "NA"
    return f"{mean:.{ndp}f} ({lo:.{ndp}f}–{hi:.{ndp}f})"


def _precision(tp: int, fp: int) -> float:
    return tp / (tp + fp) if (tp + fp) > 0 else np.nan


def _recall(tp: int, fn: int) -> float:
    return tp / (tp + fn) if (tp + fn) > 0 else np.nan


def _specificity(tn: int, fp: int) -> float:
    return tn / (tn + fp) if (tn + fp) > 0 else np.nan


def _accuracy(tp: int, fp: int, tn: int, fn: int) -> float:
    tot = tp + fp + tn + fn
    return (tp + tn) / tot if tot > 0 else np.nan


def _f1(tp: int, fp: int, fn: int) -> float:
    p = _precision(tp, fp)
    r = _recall(tp, fn)
    if np.isnan(p) or np.isnan(r) or (p + r) == 0:
        return np.nan
    return 2 * p * r / (p + r)


def _pick_best_overall_config(df_sum: pd.DataFrame) -> tuple[str, str]:
    """
    Pick one best (domain, metric) overall by averaging mean AUC across binary tasks.
    This usually returns frequency/distance for your study.
    """
    d = df_sum[(df_sum["measure"] == "auc") & (df_sum["task"].isin(TASK_ORDER))].copy()
    d["domain"] = d["domain"].astype(str).str.lower()
    d["metric"] = d["metric"].astype(str).str.lower()

    g = d.groupby(["domain", "metric"], as_index=False)["mean"].mean()
    g = g.sort_values("mean", ascending=False)
    best = g.iloc[0]
    return str(best["domain"]), str(best["metric"])


def _write_docx_table(df: pd.DataFrame, docx_path: Path, title: str) -> None:
    """
    Minimal editable DOCX writer.
    """
    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as e:
        print(f"[WARN] python-docx not available, skipping DOCX: {e}")
        return

    doc = Document()
    doc.add_paragraph(title)

    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        hdr_cells[j].text = str(col)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, col in enumerate(df.columns):
            cells[j].text = str(row[col])

    doc.save(str(docx_path))


def run(args: argparse.Namespace) -> None:
    results_dir = Path(args.results_dir).expanduser().resolve()
    out_dir = Path(args.out_dir).expanduser().resolve()
    tbl_dir = out_dir / "tables"
    tbl_dir.mkdir(parents=True, exist_ok=True)

    sum_path = results_dir / "performance_summary.csv"
    cm_path = results_dir / "confusion_matrices_aggregated.csv"
    if not sum_path.exists() or not cm_path.exists():
        raise FileNotFoundError("Missing result CSVs. Run scripts/nested_cv_train_eval.py first.")

    df_sum = pd.read_csv(sum_path)
    df_cm = pd.read_csv(cm_path)

    # -------------------------
    # Table 2: Nested CV AUC + CI95
    # -------------------------
    auc = df_sum[df_sum["measure"] == "auc"].copy()
    auc["domain"] = auc["domain"].astype(str).str.lower()
    auc["metric"] = auc["metric"].astype(str).str.lower()

    # Keep only binary tasks by default (you can expand if needed)
    auc = auc[auc["task"].isin(TASK_ORDER)].copy()

    # Build wide table rows: (Domain, Metric) x tasks
    rows = []
    for (domain, metric), g in auc.groupby(["domain", "metric"]):
        row = {
            "Domain": domain.capitalize(),
            "Metric": metric.capitalize() if metric != "distance" else "Distance",
        }
        for t in TASK_ORDER:
            gt = g[g["task"] == t]
            if gt.empty:
                row[_nice_task_label(t)] = "NA"
            else:
                r = gt.iloc[0]
                row[_nice_task_label(t)] = _fmt_auc_ci(float(r["mean"]), float(r["ci95_low"]), float(r["ci95_high"]))
        rows.append(row)

    df_table2 = pd.DataFrame(rows)
    # Sort nicely
    metric_order = {"Pearson": 0, "Spearman": 1, "Kendall": 2, "Distance": 3}
    domain_order = {"Spatial": 0, "Frequency": 1}
    df_table2["__d"] = df_table2["Domain"].map(domain_order)
    df_table2["__m"] = df_table2["Metric"].map(metric_order)
    df_table2 = df_table2.sort_values(["__d", "__m"]).drop(columns=["__d", "__m"]).reset_index(drop=True)

    out_csv2 = tbl_dir / "Table2_nested_cv_auc.csv"
    out_docx2 = tbl_dir / "Table2_nested_cv_auc.docx"
    df_table2.to_csv(out_csv2, index=False)

    _write_docx_table(
        df_table2,
        out_docx2,
        title="Table 2. Nested cross-validation discrimination performance (AUC with 95% CI) across similarity measures and domains.",
    )

    print(f"[OK] Wrote: {out_csv2}")
    print(f"[OK] Wrote: {out_docx2}")

    # -------------------------
    # Table 3: Best configuration metrics (binary tasks)
    # -------------------------
    best_domain, best_metric = _pick_best_overall_config(df_sum)

    df_best_cm = df_cm[
        (df_cm["task"].isin(TASK_ORDER))
        & (df_cm["domain"].astype(str).str.lower() == best_domain)
        & (df_cm["metric"].astype(str).str.lower() == best_metric)
    ].copy()

    if df_best_cm.empty:
        raise ValueError(f"No confusion matrices found for best config {best_domain}/{best_metric}")

    rows3 = []
    for t in TASK_ORDER:
        r = df_best_cm[df_best_cm["task"] == t]
        if r.empty:
            continue
        r = r.iloc[0]
        tp, fp, tn, fn = int(r["tp"]), int(r["fp"]), int(r["tn"]), int(r["fn"])
        acc = _accuracy(tp, fp, tn, fn)
        sens = _recall(tp, fn)
        spec = _specificity(tn, fp)
        prec = _precision(tp, fp)
        f1 = _f1(tp, fp, fn)

        rows3.append(
            {
                "Task": _nice_task_label(t),
                "Domain": best_domain.capitalize(),
                "Metric": best_metric.capitalize() if best_metric != "distance" else "Distance",
                "Positive class": str(r["positive_class"]),
                "Accuracy": f"{acc:.3f}" if not np.isnan(acc) else "NA",
                "Sensitivity": f"{sens:.3f}" if not np.isnan(sens) else "NA",
                "Specificity": f"{spec:.3f}" if not np.isnan(spec) else "NA",
                "Precision": f"{prec:.3f}" if not np.isnan(prec) else "NA",
                "F1-score": f"{f1:.3f}" if not np.isnan(f1) else "NA",
            }
        )

    df_table3 = pd.DataFrame(rows3)

    out_csv3 = tbl_dir / "Table3_best_config_metrics.csv"
    out_docx3 = tbl_dir / "Table3_best_config_metrics.docx"
    df_table3.to_csv(out_csv3, index=False)

    _write_docx_table(
        df_table3,
        out_docx3,
        title="Table 3. Operating characteristics of the best-performing configuration across binary diagnostic tasks.",
    )

    print(f"[OK] Wrote: {out_csv3}")
    print(f"[OK] Wrote: {out_docx3}")
    print(f"[INFO] Best overall config selected as: {best_domain}/{best_metric}")


def build_argparser() -> argparse.ArgumentParser:
    p = argparse.ArgumentParser(description="Build paper tables (Table 2 and Table 3).")
    p.add_argument("--results_dir", type=str, default="outputs/paper_results")
    p.add_argument("--out_dir", type=str, default="outputs/paper_outputs")
    return p


if __name__ == "__main__":
    run(build_argparser().parse_args())
